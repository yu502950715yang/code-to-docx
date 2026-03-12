import os
import sys
import docx
import argparse
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Default configuration for CLI
DEFAULT_PROJECT_NAME = ""
DEFAULT_VERSION = ""
DEFAULT_OUTPUT_FILE = f"《{DEFAULT_PROJECT_NAME}{DEFAULT_VERSION}》程序鉴别材料.docx"

# Source configurations
DEFAULT_SOURCE_CONFIGS = [
]

LINES_PER_PAGE = 50 # Estimate for 12pt line spacing + margins

def set_font_style(run, font_name='Consolas', size=10):
    """
    设置字体样式，包括中文字体支持
    """
    run.font.name = font_name
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '宋体')

def set_line_numbering(section, enable=True, count_by=1, start=1, restart='newPage'):
    """
    开启/关闭页面行号（Word 行号），支持每页重启
    """
    sect_pr = section._sectPr
    if enable:
        ln = OxmlElement('w:lnNumType')
        ln.set(qn('w:countBy'), str(count_by))
        ln.set(qn('w:start'), str(start))
        ln.set(qn('w:restart'), restart)
        # 确保只存在一个行号设置：先移除已有，再添加
        for child in list(sect_pr):
            if child.tag.endswith('lnNumType'):
                sect_pr.remove(child)
        sect_pr.append(ln)
    else:
        for child in list(sect_pr):
            if child.tag.endswith('lnNumType'):
                sect_pr.remove(child)
def suppress_line_numbers_for_paragraph(paragraph, suppress=True):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    # 移除已存在的设置
    for child in list(pPr):
        if child.tag.endswith('suppressLineNumbers'):
            pPr.remove(child)
    if suppress:
        sup = OxmlElement('w:suppressLineNumbers')
        pPr.append(sup)

def add_page_number(run):
    """
    向页脚添加页码字段
    """
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._element
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar4)

def create_header(section, text):
    """
    创建页眉并设置样式
    """
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = 'Header'
    if paragraph.runs:
        run = paragraph.runs[0]
        run.font.size = Pt(10.5) # 五号
        run.font.name = '宋体'
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), '宋体')

def create_footer(section):
    """
    创建居中的页码页脚
    """
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    add_page_number(run)

def collect_files(config):
    """
    根据配置递归收集符合扩展名的文件路径
    """
    files_found = []
    print(f"Scanning {config['name']} in {config['path']}...")
    if not os.path.exists(config['path']):
        print(f"Warning: Path does not exist: {config['path']}")
        return []
        
    for root, dirs, files in os.walk(config['path']):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in config['extensions']:
                files_found.append(os.path.join(root, file))
    return files_found

def get_all_lines(source_configs):
    """
    读取所有配置路径下的文件内容，跳过空行
    """
    all_lines = []
    for config in source_configs:
        files = collect_files(config)
        for file_path in files:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                    for line in lines:
                        if line.strip(): # Skip empty lines
                            all_lines.append(line.rstrip('\n').rstrip('\r'))
            except Exception as e:
                # print(f"Error reading {file_path}: {e}")
                pass
    return all_lines

def safe_save(doc, output_file):
    try:
        doc.save(output_file)
        return output_file
    except PermissionError:
        base, ext = os.path.splitext(output_file)
        i = 1
        while i <= 50:
            alt = f"{base}_副本{i}{ext}"
            try:
                doc.save(alt)
                print(f"Permission denied. Saved to: {alt}")
                return alt
            except PermissionError:
                i += 1
        raise

def generate_docx(max_pages, mode='sequential', project_name=DEFAULT_PROJECT_NAME, 
                  version=DEFAULT_VERSION, output_file=DEFAULT_OUTPUT_FILE, 
                  source_configs=DEFAULT_SOURCE_CONFIGS, progress_callback=None,
                  line_numbers=False):
    """
    生成 Word 文档
    """
    doc = docx.Document()
    
    # Page Setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.3)
    section.footer_distance = Cm(1.5)

    # Header & Footer
    header_text = f"{project_name}{version}"
    create_header(section, header_text)
    create_footer(section)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(200)
    title_p.paragraph_format.space_after = Pt(100)
    run = title_p.add_run(f"《{project_name}》{version}\n源代码")
    run.font.name = '黑体'
    run.font.size = Pt(26)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), '黑体')
    if line_numbers:
        suppress_line_numbers_for_paragraph(title_p, True)
    
    # 新建第二节，首页不显示行号；第二节根据开关显示行号
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    # 复制页面设置
    new_section.page_height = section.page_height
    new_section.page_width = section.page_width
    new_section.top_margin = section.top_margin
    new_section.bottom_margin = section.bottom_margin
    new_section.left_margin = section.left_margin
    new_section.right_margin = section.right_margin
    new_section.header_distance = section.header_distance
    new_section.footer_distance = section.footer_distance
    # 设置第二节的页眉页脚与行号
    create_header(new_section, header_text)
    create_footer(new_section)
    if line_numbers:
        set_line_numbering(new_section, enable=True, count_by=1, start=0, restart='newPage')

    # Content Processing
    if progress_callback: progress_callback("正在收集代码行...", 0.1)
    print("Collecting all code lines...")
    all_lines = get_all_lines(source_configs)
    total_available_lines = len(all_lines)
    print(f"Total available lines: {total_available_lines}")

    target_lines = max_pages * LINES_PER_PAGE
    lines_to_write = []

    if mode == 'start_end' and total_available_lines > target_lines:
        print(f"Mode: Start+End. Taking first {target_lines // 2} and last {target_lines // 2} lines.")
        lines_to_write.extend(all_lines[:target_lines // 2])
        lines_to_write.append("... (中间代码省略) ...")
        lines_to_write.extend(all_lines[-(target_lines // 2):])
    else:
        print(f"Mode: Sequential. Taking first {min(total_available_lines, target_lines)} lines.")
        lines_to_write = all_lines[:target_lines]

    if progress_callback: progress_callback(f"正在写入 {len(lines_to_write)} 行代码...", 0.3)
    print(f"Writing {len(lines_to_write)} lines to document...")
    
    # Write lines
    total_lines = len(lines_to_write)
    for i, line in enumerate(lines_to_write):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(12) # Exact 12pt
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        
        run = p.add_run(line)
        set_font_style(run, size=10.5) # 10.5pt (五号)
        
        if progress_callback and i % 500 == 0:
            progress = 0.3 + (i / total_lines) * 0.6
            progress_callback(f"已写入 {i}/{total_lines} 行...", progress)

    if progress_callback: progress_callback("正在保存文档...", 0.95)
    output_file = safe_save(doc, output_file)
    print(f"Document saved to: {output_file}")
    if progress_callback: progress_callback("文档生成成功！", 1.0)
    return output_file

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Generate Software Copyright Code Doc')
    parser.add_argument('--pages', type=int, default=60, help='Number of pages to generate')
    parser.add_argument('--mode', type=str, default='start_end', choices=['sequential', 'start_end'], help='Generation mode')
    parser.add_argument('--name', type=str, default=DEFAULT_PROJECT_NAME, help='Project name')
    parser.add_argument('--version', type=str, default=DEFAULT_VERSION, help='Version')
    parser.add_argument('--line-numbers', action='store_true', help='启用页面行号')
    
    args = parser.parse_args()
    
    generate_docx(args.pages, args.mode, project_name=args.name, version=args.version, line_numbers=args.line_numbers)
