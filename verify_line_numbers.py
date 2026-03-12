import os
from CodeToDocx import generate_docx
import docx

def ln_flags(path: str):
    d = docx.Document(path)
    flags = []
    for s in d.sections:
        sect_pr = s._sectPr
        flags.append(any(child.tag.endswith('lnNumType') for child in sect_pr))
    return flags
def first_paragraph_suppressed(path: str) -> bool:
    d = docx.Document(path)
    p = d.paragraphs[0]._element
    pPr = p.get_or_add_pPr()
    return any(child.tag.endswith('suppressLineNumbers') for child in pPr)

def main():
    base = os.path.dirname(__file__)
    p1 = os.path.join(base, "_tmp_line_on.docx")
    p2 = os.path.join(base, "_tmp_line_off.docx")
    print("Generating with line numbers ON...")
    generate_docx(1, mode='sequential', output_file=p1, source_configs=[], line_numbers=True)
    print("Generating with line numbers OFF...")
    generate_docx(1, mode='sequential', output_file=p2, source_configs=[], line_numbers=False)
    print(p1, ln_flags(p1))
    print(p2, ln_flags(p2))
    print("First paragraph suppressed (line_numbers ON):", first_paragraph_suppressed(p1))

if __name__ == "__main__":
    main()
